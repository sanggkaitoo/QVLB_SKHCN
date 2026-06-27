"""Kiểm tra ĐỊNH DẠNG .docx theo NĐ 30/2020 (config/format_rules.yaml).
Điểm mấu chốt: font/size có thể KẾ THỪA từ style đoạn hoặc default tài liệu,
nên phải resolve giá trị HIỆU LỰC, không chỉ đọc cấp run.
Đây là kiểm tra deterministic (bằng code) — chính xác hơn nhờ LLM.
"""
import yaml
import docx
from docx.shared import Pt, Emu

EMU_PER_CM = 360000


def load_rules(path="config/format_rules.yaml"):
    return yaml.safe_load(open(path, encoding="utf-8"))


# ---- resolve giá trị hiệu lực: run -> style -> Normal ----
def _eff_font(run, para, doc):
    if run.font.name:
        return run.font.name
    st = para.style
    while st is not None:
        if st.font and st.font.name:
            return st.font.name
        st = st.base_style
    try:
        return doc.styles["Normal"].font.name
    except Exception:
        return None


def _eff_size(run, para, doc):
    if run.font.size:
        return run.font.size.pt
    st = para.style
    while st is not None:
        if st.font and st.font.size:
            return st.font.size.pt
        st = st.base_style
    try:
        n = doc.styles["Normal"].font.size
        return n.pt if n else None
    except Exception:
        return None


def check_format(docx_path: str, rules_path="config/format_rules.yaml") -> dict:
    rules = load_rules(rules_path)
    doc = docx.Document(docx_path)
    issues = []

    # 1) Lề trang
    sec = doc.sections[0]
    for side, attr in (("top", "top_margin"), ("bottom", "bottom_margin"),
                       ("left", "left_margin"), ("right", "right_margin")):
        cm = getattr(sec, attr).cm
        rng = rules["margins_cm"][side]
        if not (rng["min"] - 0.05 <= cm <= rng["max"] + 0.05):
            issues.append({"loai": "le_trang", "vi_tri": f"lề {side}",
                           "phat_hien": f"{cm:.2f}cm",
                           "yeu_cau": f"{rng['min']}–{rng['max']}cm"})

    # 2) Font & cỡ chữ theo từng đoạn (gộp lỗi để không spam)
    allowed_fonts = set(rules["font"]["allowed"])
    allowed_sizes = set(rules["body_size_pt"]["allowed"])
    skip_empty = rules.get("skip_empty_paragraphs", True)

    bad_font, bad_size = {}, {}
    for pi, para in enumerate(doc.paragraphs):
        if skip_empty and not para.text.strip():
            continue
        for run in para.runs:
            if not run.text.strip():
                continue
            f = _eff_font(run, para, doc)
            s = _eff_size(run, para, doc)
            if f and f not in allowed_fonts:
                bad_font.setdefault(f, []).append(pi + 1)
            if s and round(s) not in allowed_sizes:
                bad_size.setdefault(round(s), []).append(pi + 1)

    for f, paras in bad_font.items():
        issues.append({"loai": "font", "phat_hien": f,
                       "yeu_cau": ", ".join(allowed_fonts),
                       "vi_tri": f"{len(paras)} đoạn (vd đoạn {paras[:5]})"})
    for s, paras in bad_size.items():
        issues.append({"loai": "co_chu", "phat_hien": f"{s}pt",
                       "yeu_cau": "/".join(map(str, allowed_sizes)) + "pt",
                       "vi_tri": f"{len(paras)} đoạn (vd đoạn {paras[:5]})"})

    # 3) Giãn dòng
    ls_max = rules["paragraph"]["line_spacing_max"]
    for pi, para in enumerate(doc.paragraphs):
        if skip_empty and not para.text.strip():
            continue
        ls = para.paragraph_format.line_spacing
        if isinstance(ls, float) and ls > ls_max + 0.01:
            issues.append({"loai": "gian_dong", "vi_tri": f"đoạn {pi+1}",
                           "phat_hien": f"{ls}", "yeu_cau": f"≤ {ls_max}"})

    return {"file": docx_path, "so_loi": len(issues), "loi": issues,
            "ket_luan": "Đạt định dạng" if not issues else "Có lỗi định dạng cần sửa"}
