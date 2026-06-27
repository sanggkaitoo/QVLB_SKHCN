"""Trích xuất text đa định dạng. Trả về (text, method).
Khác bản cũ: KHÔNG xóa file gốc; báo rõ phương pháp dùng (để lưu vào DB).
"""
import os
import subprocess
import fitz
import docx
import pandas as pd
import pytesseract
from pdf2image import convert_from_path
from PIL import Image


def extract_pdf(path: str):
    text = ""
    try:
        d = fitz.open(path)
        for page in d:
            text += page.get_text("text") + "\n"
        d.close()
    except Exception as e:
        print(f"  ! lỗi đọc PDF text {path}: {e}")
    text = text.strip()
    if len(text) >= 50:
        return text, "pdf_text"

    # PDF scan -> OCR
    print("  > PDF scan, chạy OCR tiếng Việt...")
    try:
        imgs = convert_from_path(path)
        ocr = "\n".join(pytesseract.image_to_string(im, lang="vie") for im in imgs)
        return ocr.strip(), "ocr_tesseract"
    except Exception as e:
        print(f"  ! lỗi OCR PDF: {e}")
        return "", "ocr_failed"


def extract_docx(path: str):
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:                      # lấy cả nội dung bảng
        for row in t.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts).strip(), "docx"


def extract_doc(path: str):
    tmp = "/tmp/qlvb_convert"
    os.makedirs(tmp, exist_ok=True)
    subprocess.run(["libreoffice", "--headless", "--convert-to", "docx",
                    path, "--outdir", tmp], capture_output=True)
    docx_path = os.path.join(tmp, os.path.basename(path) + "x")
    if os.path.exists(docx_path):
        text, _ = extract_docx(docx_path)
        os.remove(docx_path)
        return text, "doc_libre"
    return "", "doc_failed"


def extract_excel(path: str):
    out = ""
    xls = pd.ExcelFile(path)
    for s in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=s)
        if df.empty:
            continue
        out += f"\n--- Sheet: {s} ---\n" + df.to_csv(index=False, sep="\t") + "\n"
    return out.strip(), "xlsx"


def extract_image(path: str):
    try:
        return pytesseract.image_to_string(Image.open(path), lang="vie").strip(), "ocr_tesseract"
    except Exception as e:
        print(f"  ! lỗi OCR ảnh: {e}")
        return "", "ocr_failed"


def extract(path: str):
    ext = path.lower().rsplit(".", 1)[-1]
    try:
        if ext == "pdf":                      return extract_pdf(path)
        if ext == "docx":                     return extract_docx(path)
        if ext == "doc":                      return extract_doc(path)
        if ext in ("xlsx", "xls"):            return extract_excel(path)
        if ext in ("png", "jpg", "jpeg", "bmp", "tiff"): return extract_image(path)
    except Exception as e:
        print(f"  ! lỗi trích xuất {path}: {e}")
    return "", "unsupported"
